# --- Step 1: Imports ---
import os
import io
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
from dotenv import load_dotenv

# --- Step 2: Load OpenAI API key ---
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# --- Step 3: Load Excel file ---
file_path = "AccidentDatabase.xlsx"  # 🔁 Replace with your file path
df = pd.read_excel(file_path)
print("✅ Excel file loaded successfully.")

# --- Step 4: Identify useful categorical columns ---
excluded_cols = ['Latitude', 'Longitude', 'X-Coordinate', 'Y-Coordinate']
categorical_cols = [
    col for col in df.columns
    if df[col].dtype == 'object'
    and col not in excluded_cols
    and 2 <= df[col].nunique() <= 15
]
print("📊 Selected columns for analysis:", categorical_cols)

# --- Step 5: Create a new Word document ---
doc = Document()
doc.add_heading("Collision Analysis Report", 0)
doc.add_paragraph("Generated automatically by Mobility Edge Solutions")
doc.add_page_break()

figure_count = 1

# --- Step 6: Generate sections for each column ---
for col in categorical_cols:
    print(f"\n🧩 Processing column: {col}")
    value_counts = df[col].value_counts()

    if len(value_counts) < 2:
        continue

    # --- Step 6.1: Create chart in memory ---
    img_stream = io.BytesIO()
    plt.figure(figsize=(6, 4))
    if len(value_counts) <= 5:
        plt.pie(value_counts, labels=value_counts.index, autopct='%1.1f%%')
    else:
        value_counts.plot(kind='bar')
        plt.xticks(rotation=45, ha='right')
    plt.title(col)
    plt.tight_layout()
    plt.savefig(img_stream, format='png')
    plt.close()
    img_stream.seek(0)

    # --- Step 6.2: Generate GPT summary using your prompt ---
    try:
        title = col
        data = value_counts.to_string()
        prompt = (
            f"You are a road safety analyst. Write a short professional summary of the chart titled '{title}'.\n\n"
            f"Data: {data}\n\n"
            f"Highlight notable patterns, especially frequencies, dominant values, or changes over time. "
            f"Use a tone similar to a traffic safety expert / consultant writing for a municipality."
        )

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=250
        )
        summary = response.choices[0].message.content.strip()
    except Exception as e:
        summary = f"[⚠️ GPT error: {e}]"
        print(summary)

    # --- Step 6.3: Insert into Word report ---
    doc.add_heading(f"{figure_count}. {col}", level=1)
    doc.add_picture(img_stream, width=Inches(5.5))

    caption = doc.add_paragraph(f"Figure {figure_count}: {col} Distribution")
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption.runs[0].italic = True

    paragraph = doc.add_paragraph(summary)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.runs[0].font.size = Pt(11)

    doc.add_page_break()
    figure_count += 1

# --- Step 7: Save the report ---
output_file = "collision_report.docx"
doc.save(output_file)
print(f"\n✅ Report saved successfully as '{output_file}'")
