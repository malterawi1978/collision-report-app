import os
import io
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
from datetime import datetime
from PIL import Image
import geopandas as gpd
from shapely.geometry import Point
import contextily as ctx
from matplotlib.patches import FancyArrow

st.set_page_config(
    page_title="Collisio – Collision Report Generator",
    page_icon="🚦",
    layout="centered"
)

logo = Image.open("Collisio_Logo.png")
st.image(logo, width=100)

st.title("🤖 Collisio")
st.markdown("### Collision Report Generator")
st.markdown("Upload your traffic accident data to generate a smart report with charts and insights powered by AI.")

st.markdown("**Need help formatting your accident data?**")
st.markdown("Download our ready-made Excel template to ensure your data is structured correctly before upload.")

with open("collision_template.xlsx", "rb") as f:
    st.download_button(
        label="Download Excel Template",
        data=f,
        file_name="collision_template.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

uploaded_file = st.file_uploader("📂 Upload Excel File", type=["xlsx"])

if uploaded_file:
    progress = st.progress(0, text="Starting report generation...")
        left_col, right_col = st.columns([4, 1])
        with left_col:
            pass
        with right_col:
            st.markdown("### 📋 Sections")
            section_placeholder = st.empty()
        steps = 20
        current_step = 0
        def update_progress(msg):
            nonlocal current_step
            current_step += 1
            progress.progress(min(current_step, steps) / steps, text=msg)
        df = pd.read_excel(uploaded_file)
        df.dropna(how='all', inplace=True)
        df = df.dropna(subset=['Classification Of Accident'])
        df = df.applymap(lambda x: x.strip().replace("**", "").replace("###", "") if isinstance(x, str) else x)
        df = df[~df.isin(['', ' ', None]).any(axis=1)]
        update_progress("Excel file read successfully.")
        st.success("File read successfully.")

        doc = Document()
        doc.add_heading("Collision Analysis Report", 0)
        doc.add_paragraph("Prepared by Mobility Edge Solution").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_page_break()

        section_count = 1

        def add_section(title, chart_data, chart_type="bar"):
            global section_count
            if chart_data.empty:
                return

            img_stream = io.BytesIO()
            plt.figure(figsize=(6, 4))
            if chart_type == "pie" and len(chart_data) <= 6:
                plt.pie(chart_data, labels=chart_data.index, autopct='%1.1f%%')
            else:
                ax = chart_data.plot(kind="bar", stacked=isinstance(chart_data, pd.DataFrame))
                plt.xticks(rotation=45, ha='right')
                ax.set_ylabel("Number of Accidents")
            plt.title(title)
            plt.tight_layout()
            plt.savefig(img_stream, format="png")
            plt.close()
            img_stream.seek(0)

            prompt = (
                f"You are a road safety analyst. Write a short professional summary "
                f"of this accident chart titled '{title}'.\n\n"
                f"Data: {chart_data.head(10).to_string()}\n\n"
                f"Highlight the most common types and any interesting patterns."
            )
            try:
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=300
                )
                summary = response.choices[0].message.content.strip()
            except Exception as e:
                summary = f"[GPT Error: {e}]"

            clean_title = title.replace("**", "").replace("###", "").strip().replace("#", "").strip()
            section_placeholder.markdown(f"**Section {section_count}: {clean_title}**")
            update_progress(f"Analyzing {clean_title}")
            doc.add_heading(f"Section {section_count}: {clean_title}", level=1)
            doc.add_picture(img_stream, width=Inches(5.5))
            caption = doc.add_paragraph(f"Figure {section_count}: {clean_title}")
            caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            caption.runs[0].italic = True
            clean_summary = summary.replace("**", "").replace("###", "").strip()
            para = doc.add_paragraph(clean_summary)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            para.runs[0].font.size = Pt(11)
            doc.add_page_break()
            section_count += 1

        def add_grouped_section(column_name, title):
            if column_name in df.columns and 'Classification Of Accident' in df.columns:
                grouped = df.groupby([column_name, 'Classification Of Accident']).size().unstack(fill_value=0)
                if not grouped.empty:
                    add_section(title, grouped, chart_type="bar")

        if 'Classification Of Accident' in df.columns:
            add_section("Accident Severity Distribution", df['Classification Of Accident'].value_counts(), chart_type="pie")

        add_grouped_section("Accident Year", "Accidents by Year")
        add_grouped_section("Accident Day", "Accidents by Day")
        add_grouped_section("Light", "Light Condition Distribution")
        add_grouped_section("Environment Condition 1", "Environment Condition 1 Distribution")
        add_grouped_section("Environment Condition 2", "Environment Condition 2 Distribution")
        add_grouped_section("Initial Impact Type", "Initial Impact Type Analysis")
        add_grouped_section("Impact Location", "Impact Location Analysis")
        add_grouped_section("Apparent Driver 1 Action", "Driver 1 Action Trends")
        add_grouped_section("Apparent Driver 2 Action", "Driver 2 Action Trends")
        add_grouped_section("Driver 1 Condition", "Driver 1 Condition Trends")
        add_grouped_section("Driver 2 Condition", "Driver 2 Condition Trends")

        if 'Accident Date' in df.columns and 'Classification Of Accident' in df.columns:
            try:
                df['Accident Date'] = pd.to_datetime(df['Accident Date'], errors='coerce')
                df['Day of Week'] = df['Accident Date'].dt.day_name()
                weekday_order = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
                weekday_grouped = df.groupby(['Day of Week', 'Classification Of Accident']).size().unstack(fill_value=0).reindex(weekday_order)
                add_section("Accident Type by Day of Week", weekday_grouped, chart_type="bar")

                df['Day Type'] = df['Accident Date'].dt.dayofweek.apply(lambda x: 'Weekend' if x >= 5 else 'Weekday')
                daytype_grouped = df.groupby(['Day Type', 'Classification Of Accident']).size().unstack(fill_value=0)
                add_section("Accident Type by Weekday vs Weekend", daytype_grouped, chart_type="bar")

                df['Accident Month'] = df['Accident Date'].dt.month_name()
                month_order = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']
                month_grouped = df.groupby(['Accident Month', 'Classification Of Accident']).size().unstack(fill_value=0).reindex(month_order)
                add_section("Accident Type by Month", month_grouped, chart_type="bar")
            except Exception as e:
                st.warning(f"Could not process date-based charts: {e}")

        if 'Accident Time' in df.columns and 'Classification Of Accident' in df.columns:
            try:
                def clean_time_string(t):
                    if pd.isnull(t): return None
                    t = str(t).lower().strip().replace('pm', '').replace('am', '')
                    try:
                        dt = datetime.strptime(t.strip(), '%H:%M:%S')
                    except:
                        try:
                            dt = datetime.strptime(t.strip(), '%I:%M:%S')
                        except:
                            return None
                    return dt.strftime('%H:%M')

                df['Cleaned Time'] = df['Accident Time'].apply(clean_time_string)

                def classify_period(t):
                    try:
                        if t is None:
                            return 'Unknown'
                        hour = int(t.split(':')[0])
                        if 6 <= hour < 12: return 'Morning'
                        elif 12 <= hour < 17: return 'Afternoon'
                        elif 17 <= hour < 21: return 'Evening'
                        else: return 'Night'
                    except:
                        return 'Unknown'

                df['Time Period'] = df['Cleaned Time'].apply(classify_period)
                period_order = ['Morning', 'Afternoon', 'Evening', 'Night']
                df_filtered = df[df['Time Period'].isin(period_order)]
                time_grouped = df_filtered.groupby(['Time Period', 'Classification Of Accident']).size().unstack(fill_value=0).reindex(period_order)
                add_section("Accident Type by Time of Day", time_grouped, chart_type="bar")
            except Exception as e:
                st.warning(f"Could not process time of day: {e}")

        try:
            if 'Latitude' in df.columns and 'Longitude' in df.columns and 'Classification Of Accident' in df.columns:
                df["Classification Of Accident"] = df["Classification Of Accident"].astype(str).str.strip().str.lower()
                color_list = [
                    '#FF0000', '#00CC00', '#0000FF', '#FFA500', '#800080',
                    '#00FFFF', '#FFC0CB', '#FFFF00', '#00CED1', '#FF1493'
                ]
                unique_types = df["Classification Of Accident"].unique()
                auto_color_map = {stype: color_list[i % len(color_list)] for i, stype in enumerate(unique_types)}

                geometry = [Point(xy) for xy in zip(df["Longitude"], df["Latitude"])]
                gdf = gpd.GeoDataFrame(df, geometry=geometry, crs="EPSG:4326").to_crs(epsg=3857)

                fig, ax = plt.subplots(figsize=(14, 12))
                for acc_type in unique_types:
                    subset = gdf[gdf["Classification Of Accident"] == acc_type]
                    color = auto_color_map[acc_type]
                    label = acc_type.title()
                    subset.plot(ax=ax, label=label, color=color, markersize=100, edgecolor='none')

                buffer = 500
                minx, miny, maxx, maxy = gdf.total_bounds
                ax.set_xlim(minx - buffer, maxx + buffer)
                ax.set_ylim(miny - buffer, maxy + buffer)
                ctx.add_basemap(ax, source=ctx.providers.OpenStreetMap.Mapnik, zoom=17)

                ax.text(0.95, 0.95, 'N', transform=ax.transAxes,
                        fontsize=20, fontweight='bold', ha='center', va='center', color='black')
                arrow = FancyArrow(0.95, 0.91, 0, 0.03, transform=ax.transAxes,
                                   width=0.01, head_width=0.03, head_length=0.02,
                                   length_includes_head=True, color='black', edgecolor='white')
                ax.add_patch(arrow)

                ax.set_title("Accident Locations by Type", fontsize=16)
                ax.axis("off")
                ax.legend(title="Accident Type", fontsize=10, title_fontsize=11, loc="lower left")

                plt.tight_layout()
                map_path = "accident_map.png"
                plt.savefig(map_path, dpi=600)
                plt.close()
        except Exception as e:
            st.warning(f"Could not generate street map: {e}")

        doc.add_heading(f"Section {section_count}: Collision Type Diagrams", level=1)
        doc.add_paragraph("[Custom collision type diagrams will be rendered based on type and geometry data in future versions.]")
        doc.add_page_break()

        output_path = "collision_report.docx"
        update_progress("Finalizing and saving the report")
        doc.save(output_path)
        st.success("✅ Report is ready!")
        with open(output_path, "rb") as f:
            st.download_button("🧾 Download Report", f, file_name="collision_report.docx")

        with open(map_path, "rb") as img_file:
            st.markdown("**🗺️ Download Accident Map**")
            st.download_button("🗺️ Download Map (PNG)", img_file.read(), file_name="accident_map.png", mime="image/png")
else:
    st.info("Please upload an Excel file to begin.")
